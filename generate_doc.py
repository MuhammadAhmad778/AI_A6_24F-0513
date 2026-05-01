from docx import Document

doc = Document()
doc.add_heading('AI2002 - Artificial Intelligence', 0)
doc.add_heading('Assignment 06: Question 6 Project Report', 1)
doc.add_paragraph('Name: Muhammad Ahmad')
doc.add_paragraph('Roll Number: 24F-0513')

doc.add_heading('Project Links', 2)
doc.add_paragraph('GitHub Repository: https://github.com/MuhammadAhmad778/AI_A6_24F-0513')
doc.add_paragraph('Live Vercel Deployment: https://ai-a6-24-f-0513.vercel.app/')
doc.add_paragraph('LinkedIn Demo Post: [INSERT_YOUR_LINKEDIN_POST_LINK_HERE]')

doc.add_heading('1. Project Overview', 1)
doc.add_paragraph('For this project, I developed a Dynamic Wumpus Logic Agent accessible via a modern web interface. The system comprises a React-based frontend built with Next.js and a Python inference engine built with FastAPI.')
doc.add_paragraph('Unlike traditional pathfinding algorithms (like A* or BFS), this agent operates purely on Propositional Logic. It navigates the grid safely by receiving real-time percepts (Breezes and Stenches) and deducing the safety of adjacent cells using a formal Knowledge Base.')

doc.add_heading('2. The Logic Pipeline', 1)
doc.add_heading('A. The Knowledge Base (KB)', 2)
doc.add_paragraph('The agent\'s "brain" is its Knowledge Base. It starts with basic facts (e.g., the starting cell [0, 0] is safe). As the agent moves, it adds new logical rules based on the percepts it encounters.')
doc.add_paragraph('For example, if the agent moves to a cell and feels a Breeze, it knows at least one adjacent cell contains a Pit. It formalizes this as:\nB_x_y <=> (P_n1 V P_n2 V P_n3 ...)')

doc.add_heading('B. Conjunctive Normal Form (CNF) Conversion', 2)
doc.add_paragraph('To allow our Python backend to automatically prove theorems, all rules entering the KB must be converted into Conjunctive Normal Form (CNF). A sentence in CNF is a series of ANDs, where each element is an OR of literals.')
doc.add_paragraph('Conversion Steps used in the engine:\n1. Eliminate Implications (=>): A => B becomes ¬A V B.\n2. Eliminate Biconditionals (<=>): A <=> B becomes (¬A V B) AND (¬B V A).\n3. Move Negations Inward: Using De Morgan\'s Laws, e.g., ¬(A AND B) becomes ¬A V ¬B.\n4. Distribute V over AND: To ensure the final structure is strictly clauses joined by ANDs.')
doc.add_paragraph('In my Python implementation, these CNF clauses are stored highly efficiently as sets of strings within a frozenset. For example, (¬P_1_1 V B_2_1) is stored as frozenset(["-P_1_1", "B_2_1"]).')

doc.add_heading('3. The Resolution Refutation Loop', 1)
doc.add_paragraph('Before the agent enters any unknown adjacent cell, it must mathematically prove that the cell is safe. It asks two questions:\n1. "Is there a pit here?" (Goal: -P_x_y)\n2. "Is there a Wumpus here?" (Goal: -W_x_y)')
doc.add_paragraph('To answer these, the backend utilizes the Resolution Refutation algorithm.')

doc.add_heading('The Algorithm: Proof by Contradiction', 2)
doc.add_paragraph('1. Negate the Goal: If the agent wants to prove -P_x_y (the cell is safe from a pit), it temporarily assumes the opposite: P_x_y (there IS a pit). It adds this negated goal to a temporary copy of the KB.')
doc.add_paragraph('2. Resolve Clauses: The engine iterates through all pairs of clauses in the KB. It looks for complementary literals (e.g., P_1_1 in one clause and -P_1_1 in another).')
doc.add_paragraph('3. Generate Resolvents: When a complementary pair is found, the literals cancel out, and the remaining literals from both clauses are merged into a new, smaller clause.')
doc.add_paragraph('4. Find the Empty Clause (Contradiction): If resolving two unit clauses (e.g., P_1_1 and -P_1_1) results in an empty clause (contradiction), a logical contradiction has occurred.')
doc.add_paragraph('5. Conclusion: Because assuming P_x_y led to a contradiction, P_x_y must be false. Therefore, the original goal (-P_x_y) is strictly entailed by the Knowledge Base. The cell is 100% safe to enter.')

doc.add_heading('Algorithmic Efficiency', 2)
doc.add_paragraph('To prevent infinite loops during resolution (which can occur in highly complex logic bases), the engine utilizes a bounded search. It limits the number of pairs checked per iteration and breaks early if no new resolvents are generated, ensuring real-time responsiveness for the web application.')

doc.save('AI_A6_24F-0513_Report.docx')
print('Word document created successfully!')
